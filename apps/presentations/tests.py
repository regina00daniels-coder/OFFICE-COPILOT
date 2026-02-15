from io import BytesIO

from docx import Document
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import Client, TestCase
from django.urls import reverse

from apps.tenants.models import Tenant
from .models import Presentation


class WordToPresentationTests(TestCase):
    def setUp(self):
        self.tenant = Tenant.objects.create(name="Tenant A", domain="a.local")
        user_model = get_user_model()
        self.user = user_model.objects.create_user(
            username="writer",
            password="pass1234",
            tenant=self.tenant,
            role=user_model.Role.ADMIN,
        )
        self.client = Client()

    def _build_docx_upload(self):
        document = Document()
        document.add_heading("Q1 Operations Review", level=1)
        document.add_paragraph("Revenue grew by 14 percent across enterprise clients.")
        document.add_paragraph("Hiring pipeline remains constrained in technical roles.")
        document.add_heading("Actions", level=1)
        document.add_paragraph("Expand partner channel and reduce onboarding delays.")
        output = BytesIO()
        document.save(output)
        return SimpleUploadedFile(
            "report.docx",
            output.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def test_word_to_presentation_api_creates_presentation(self):
        self.client.login(username="writer", password="pass1234")
        response = self.client.post(
            reverse("word-to-presentation"),
            data={"file": self._build_docx_upload()},
            HTTP_X_TENANT="a.local",
            HTTP_HOST="localhost",
        )
        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertIn("presentation_id", payload)
        presentation = Presentation.objects.get(id=payload["presentation_id"])
        self.assertEqual(presentation.tenant, self.tenant)
        self.assertTrue(bool(presentation.file))
