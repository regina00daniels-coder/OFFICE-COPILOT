import json
from io import BytesIO
from unittest.mock import patch

from docx import Document
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import Client, TestCase
from django.urls import reverse
from openpyxl import Workbook

from apps.tenants.models import Tenant
from .models import DataAnalysisRun, DocumentReportRun
from .services import analyze_business_data


class ReportingRoleAccessTests(TestCase):
    def setUp(self):
        tenant = Tenant.objects.create(name="Tenant A", domain="a.local")
        user_model = get_user_model()
        self.staff = user_model.objects.create_user(
            username="staff", password="pass1234", tenant=tenant, role=user_model.Role.STAFF
        )
        self.member = user_model.objects.create_user(
            username="member", password="pass1234", tenant=tenant, role=user_model.Role.USER
        )
        self.client = Client()

    def test_staff_can_create_report(self):
        self.client.login(username="staff", password="pass1234")
        response = self.client.post(
            reverse("report-list-create"),
            data=json.dumps({"name": "Ops Weekly"}),
            content_type="application/json",
            HTTP_X_TENANT="a.local",
            HTTP_HOST="localhost",
        )
        self.assertEqual(response.status_code, 201)

    def test_regular_user_cannot_create_report(self):
        self.client.login(username="member", password="pass1234")
        response = self.client.post(
            reverse("report-list-create"),
            data=json.dumps({"name": "Ops Weekly"}),
            content_type="application/json",
            HTTP_X_TENANT="a.local",
            HTTP_HOST="localhost",
        )
        self.assertEqual(response.status_code, 403)

    def _dataset_upload(self):
        workbook = Workbook()
        ws = workbook.active
        ws.append(["department", "revenue", "cost", "region"])
        ws.append(["Surgery", 120000, 70000, "North"])
        ws.append(["Surgery", 128000, 71000, "North"])
        ws.append(["Pharmacy", 85000, 35000, "East"])
        ws.append(["Logistics", 64000, 28000, "West"])
        payload = BytesIO()
        workbook.save(payload)
        return SimpleUploadedFile(
            "hospital_data.xlsx",
            payload.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    def _doc_upload(self):
        doc = Document()
        doc.add_heading("Weekly Operations Brief", level=1)
        doc.add_paragraph("Hospital inventory remains within safety stock thresholds.")
        doc.add_paragraph("Outpatient visits increased by 9 percent week over week.")
        content = BytesIO()
        doc.save(content)
        return SimpleUploadedFile(
            "brief.docx",
            content.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def test_staff_can_run_data_analysis_workflow(self):
        self.client.login(username="staff", password="pass1234")
        response = self.client.post(
            reverse("reporting-data-run"),
            data={"file": self._dataset_upload()},
            HTTP_X_TENANT="a.local",
            HTTP_HOST="localhost",
        )
        self.assertEqual(response.status_code, 302)
        run = DataAnalysisRun.objects.first()
        self.assertIsNotNone(run)
        self.assertEqual(run.status, DataAnalysisRun.Status.COMPLETED)
        self.assertTrue(bool(run.workbook_file))

    def test_staff_can_generate_document_powerpoint_report(self):
        self.client.login(username="staff", password="pass1234")
        response = self.client.post(
            reverse("reporting-doc-run"),
            data={"file": self._doc_upload()},
            HTTP_X_TENANT="a.local",
            HTTP_HOST="localhost",
        )
        self.assertEqual(response.status_code, 302)
        run = DocumentReportRun.objects.first()
        self.assertIsNotNone(run)
        self.assertEqual(run.status, DocumentReportRun.Status.COMPLETED)
        self.assertTrue(bool(run.powerpoint_file))

    def test_large_dataset_is_split_across_multiple_cleaned_data_sheets(self):
        workbook = Workbook()
        ws = workbook.active
        ws.append(["department", "revenue"])
        for i in range(45):
            ws.append([f"Dept{i % 3}", i * 10])
        content = BytesIO()
        workbook.save(content)
        content.seek(0)
        with patch("apps.reporting.services.MAX_DATA_ROWS_PER_SHEET", 10):
            summary, _ = analyze_business_data(content, "large_data.xlsx")
        self.assertGreaterEqual(summary["cleaned_data_sheets"], 5)
        self.assertIn("cleaned_rows_exported", summary)
